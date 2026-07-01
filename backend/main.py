import os
import random
import json
import zipfile
import threading
import uuid
from io import BytesIO
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed

from fastapi import FastAPI, Response, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from starlette.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import create_engine, Column, Integer, String, JSON, DateTime, Text
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker
from google import genai
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(title="Role 2 Exercise Builder API")

# Database setup
DATABASE_URL = os.getenv("DATABASE_URL")
if DATABASE_URL and DATABASE_URL.startswith("postgres://"):
    DATABASE_URL = DATABASE_URL.replace("postgres://", "postgresql://", 1)

engine = create_engine(DATABASE_URL) if DATABASE_URL else None
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine) if engine else None
Base = declarative_base()

class Exercise(Base):
    __tablename__ = "exercises"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String, index=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    config = Column(JSON)
    cases = Column(JSON)
    msel_data = Column(JSON)
    warno_text = Column(Text)
    annex_q_text = Column(Text)
    medroe_text = Column(Text)
    road_to_war_text = Column(Text)

class Job(Base):
    __tablename__ = "jobs"
    id = Column(String, primary_key=True)
    status = Column(String, default="running")
    progress = Column(String, default="Starting...")
    completed = Column(Integer, default=0)
    total = Column(Integer, default=0)
    token = Column(String, nullable=True)
    filename = Column(String, nullable=True)
    error = Column(Text, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)

if engine:
    try:
        Base.metadata.create_all(bind=engine)
    except Exception as _e:
        print(f"WARNING: DB table creation failed: {_e}")
    # create_all does not ALTER existing tables — add newer columns idempotently.
    try:
        from sqlalchemy import text as _sql_text
        with engine.begin() as _conn:
            _conn.execute(_sql_text(
                "ALTER TABLE exercises ADD COLUMN IF NOT EXISTS road_to_war_text TEXT"
            ))
    except Exception as _e:
        print(f"WARNING: DB column migration (road_to_war_text) failed: {_e}")

CORS_ORIGINS = os.getenv("CORS_ORIGINS", "*").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_client = None

# Gemini model id. Centralized + env-overridable so a Google model retirement
# (e.g. gemini-2.0-flash was decommissioned with a 404) is a one-variable fix
# on Railway rather than a code change + redeploy.
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash")

def get_client():
    global _client
    if _client is None:
        key = os.getenv("GEMINI_API_KEY")
        if not key:
            raise HTTPException(status_code=503, detail="GEMINI_API_KEY not set on server")
        _client = genai.Client(api_key=key)
    return _client

def _response_text(response) -> str:
    """Safely pull text out of a Gemini response.

    response.text raises when the model returns no usable candidate (e.g. the
    prompt tripped a safety filter). Fall back to walking the candidate parts,
    and return "" rather than raising so callers can degrade gracefully.
    """
    try:
        text = response.text
        if text:
            return text.strip()
    except Exception:
        pass
    try:
        parts = []
        for cand in getattr(response, "candidates", None) or []:
            content = getattr(cand, "content", None)
            for part in (getattr(content, "parts", None) or []):
                if getattr(part, "text", None):
                    parts.append(part.text)
        return "".join(parts).strip()
    except Exception:
        return ""

# Curated pools mirror the /generate-name prompt so the offline fallback stays
# on-brand. None of these appear in the prompt's banned-word list.
_NAME_DESCRIPTORS = [
    "Silent", "Relentless", "Fractured", "Bleeding", "Stalking", "Burning",
    "Severed", "Hollow", "Bitter", "Fading", "Broken", "Restless", "Savage",
    "Grim", "Fallen", "Rising", "Distant", "Shattered", "Quiet", "Vigilant",
]
_NAME_NOUNS_LAND = [
    "Jackal", "Wolverine", "Lynx", "Mantis", "Fury", "Wrath", "Valor",
    "Resolve", "Defiance", "Reckoning", "Rampart", "Bastion", "Crucible",
    "Gauntlet", "Threshold", "Salient", "Bulwark", "Sentinel", "Vanguard",
]
_NAME_NOUNS_MARITIME = ["Mako", "Condor", "Viper", "Barracuda", "Kraken", "Tempest"]

def _local_exercise_name(region: str = "") -> str:
    """Deterministic-random offline name so the endpoint never fails to fetch."""
    nouns = list(_NAME_NOUNS_LAND)
    if any(w in (region or "").lower() for w in ("marit", "pacific", "naval", "littoral", "coast", "sea")):
        nouns = nouns + _NAME_NOUNS_MARITIME
    return f"Operation {random.choice(_NAME_DESCRIPTORS)} {random.choice(nouns)}"

# Short-lived store for completed zip packages keyed by download token
_packages: Dict[str, bytes] = {}

# Pydantic Models
class DayConfig(BaseModel):
    day_number: int
    tactical_setting: str
    total_patients: int
    total_waves: int
    night_ops: bool = False
    mascal: bool = False
    mascal_etiology: Optional[str] = None
    mascal_patients: Optional[int] = None
    cbrn: bool = False
    detainee_ops: bool = False

class ExerciseConfig(BaseModel):
    exercise_name: str
    duration: int
    supported_unit: str
    environment: str
    threat_level: str
    region: str
    selected_mets: List[str]
    selected_footprint: List[str]
    specialists: Dict[str, int]
    days: List[DayConfig]

# Case data
DNBI_BY_ENVIRONMENT = {
    "Jungle": ["Dengue hemorrhagic fever", "Malaria", "Snake envenomation", "Cellulitis", "Heat exhaustion", "Leptospirosis"],
    "Desert": ["Heat stroke", "Severe dehydration", "Scorpion envenomation", "Sand fly fever", "Rhabdomyolysis"],
    "Arctic": ["Severe frostbite", "Hypothermia", "Trench foot", "Cold urticaria"],
    "Maritime": ["Near-drowning", "Decompression sickness", "Marine envenomation", "Saltwater aspiration"],
    "Urban": ["Crush injury", "Smoke inhalation", "MSK overuse", "CO poisoning"],
    "Mountain": ["HAPE", "HACE", "Acute mountain sickness", "Fall with fractures"],
    "Littoral": ["Near-drowning", "Jellyfish envenomation", "Coral cuts with infection", "Heat exhaustion"],
    "General": ["Combat stress reaction", "Dental emergency", "Acute appendicitis", "Kidney stones", "Pneumonia", "Gastroenteritis"]
}

TRAUMA_BY_ETIOLOGY = {
    "IED/Blast": ["Traumatic bilateral lower extremity amputation", "TBI with penetrating fragment", "Blast lung injury", "Penetrating abdominal trauma with evisceration", "Severe burns with blast injury"],
    "Vehicle Rollover": ["Blunt abdominal trauma with splenic laceration", "C-spine fracture", "Bilateral femur fractures", "Flail chest", "Pelvic fracture with hemorrhage"],
    "GSW/Small Arms": ["Penetrating chest trauma with hemothorax", "GSW abdomen with liver laceration", "GSW extremity with vascular injury", "GSW neck with airway compromise"],
    "Aviation Mishap": ["Multi-system blunt trauma", "Severe burns (>40% TBSA)", "Spinal cord injury", "Traumatic brain injury"],
    "Indirect Fire/Mortar": ["Multiple penetrating fragment wounds", "TBI from blast", "Traumatic amputation", "Penetrating eye injury"],
    "Structural Collapse": ["Crush syndrome", "Compartment syndrome", "Traumatic asphyxiation", "Multiple long bone fractures"],
    "Burns/Fire": ["Severe thermal burns (>30% TBSA)", "Inhalation injury", "CO poisoning", "Facial burns with airway compromise"],
    "Drowning (Amphibious)": ["Near-drowning with aspiration", "Hypothermia with near-drowning", "Trauma from vessel impact"],
    "VBIED": ["Multi-system blast trauma", "Severe burns with TBI", "Traumatic amputation bilateral", "Penetrating torso wounds"],
}

GENERAL_TRAUMA = [
    "GSW right thigh with femoral artery injury", "Penetrating abdominal trauma", "Blunt chest trauma with rib fractures",
    "Open tibia-fibula fracture", "Traumatic brain injury - moderate", "Blast injury with TM rupture",
    "Penetrating neck trauma", "GSW chest with pneumothorax", "Pelvic fracture - stable", "Burn injury 15% TBSA"
]

def determine_case_phases(case_type: str, mechanism: str, is_mascal: bool = False) -> List[str]:
    surgical_kw = [
        "amputation", "evisceration", "vascular injury", "hemothorax", "fasciotomy",
        "thoracotomy", "GSW abdomen", "penetrating chest", "penetrating abdominal",
        "pelvic fracture with hemorrhage", "compartment syndrome", "crush syndrome",
        "burns (>30%", "burns (>40%", "severe thermal burns",
        "liver laceration", "splenic laceration", "bilateral femur",
        "penetrating torso", "multi-system blast trauma", "airway compromise",
        "long bone fracture",
    ]
    non_surgical_kw = [
        # Environmental / DNBI
        "dengue", "malaria", "fever", "heat stroke", "hypothermia", "altitude",
        "HAPE", "HACE", "pneumonia", "gastroenteritis", "appendicitis",
        "kidney stones", "combat stress", "dental", "rhabdomyolysis",
        "decompression sickness", "sand fly", "leptospirosis", "cold urticaria",
        # Neuro / head
        "TBI", "concussion", "closed head", "traumatic brain injury",
        # Pulmonary / inhalation (medical management, not DCS)
        "blast lung", "inhalation injury", "CO poisoning", "smoke inhalation",
        "near-drowning", "drowning", "aspiration", "saltwater aspiration",
        # Envenomation / marine
        "envenomation", "jellyfish", "coral",
        # Spinal (stabilisation, not immediate DCS)
        "spinal cord", "C-spine",
    ]

    case_lower = (case_type + " " + mechanism).lower()

    for kw in surgical_kw:
        if kw.lower() in case_lower:
            return ["DCR", "DCS", "PCC"]
    for kw in non_surgical_kw:
        if kw.lower() in case_lower:
            return ["DCR", "PCC"]

    # MASCAL waves produce high-acuity trauma — unmatched defaults to surgical.
    # Routine waves are mixed-acuity — unmatched defaults to non-surgical.
    return ["DCR", "DCS", "PCC"] if is_mascal else ["DCR", "PCC"]

CASE_SYSTEM_PROMPT = """You are an expert Military Medical Simulation Designer for Role 2 (Forward Surgical) care.
Generate a detailed simulation case with:
1. Z-MIST (Zap=EXACTLY 5 digits, Mechanism, Injuries, Signs, Treatment)
2. 9-Line Medevac (NATO format)
3. Clinical Phases (only include phases specified - some cases don't need surgery)
4. Vitals trends (3-4 timestamps per phase)
5. Contingencies (If/Then decision points)
6. Evacuation (transport type, considerations, handover notes)
7. Learning Objectives (3-5 specific, measurable)
8. Debrief Questions (4-6 thought-provoking)

If DCS not needed, set dcs to null.
Set triage_category to "Expectant" only for injuries not survivable with the
resources at hand; use it sparingly. Set disposition to the realistic outcome
after Role 2 care for THIS casualty (RTD for minor, Evac to Role 3 for surgical/
serious, Hold at Role 2 for short observation, Died of Wounds or Expectant for
non-survivable). Do NOT include blood unit counts — those are computed by a
planning factor, not per case.

JSON STRUCTURE:
{
  "meta": {"title": "String", "estimated_duration": "String", "personnel": "String", "target_specialty": "String"},
  "learning_objectives": ["String"],
  "zmist": {"zap": "USE EXACT VALUE PROVIDED IN PROMPT", "mechanism": "String", "injuries": "String", "signs": "String", "treatment": "String"},
  "nine_line": {"line1_location": "String", "line2_freq": "String", "line3_patients_precedence": "String", "line4_equipment": "String", "line5_patients_type": "String", "line6_security": "String", "line7_marking": "String", "line8_nationality": "String", "line9_nbc_terrain": "String"},
  "patient_data": {"demographics": "String", "history": "String", "allergies": "String"},
  "triage_category": "T1/T2/T3/T4/Expectant",
  "disposition": "One of: RTD | Evac to Role 3 | Hold at Role 2 | Died of Wounds | Expectant",
  "phases": {
    "dcr": {"title": "Damage Control Resuscitation", "narrative": "String", "expected_actions": ["String"], "vitals_trend": [{"time": "String", "hr": "String", "bp": "String", "rr": "String", "spo2": "String", "gcs": "String"}], "contingencies": [{"condition": "String", "consequence": "String", "intervention": "String"}]},
    "dcs": null or {"title": "Damage Control Surgery", "narrative": "String", "expected_actions": ["String"], "vitals_trend": [...], "contingencies": [...]},
    "pcc": {"title": "Prolonged Casualty Care", "narrative": "String", "expected_actions": ["String"], "vitals_trend": [...], "contingencies": [...]}
  },
  "labs": {"hgb": "String", "ph": "String", "lactate": "String", "base_excess": "String", "inr": "String"},
  "evacuation": {"transport_type": "String", "priority": "String", "considerations": "String", "handover_notes": "String"},
  "debrief_questions": ["String"]
}"""

def generate_case_sync(case_type: str, mechanism: str, environment: str, region: str, is_mascal: bool = False) -> Dict:
    phases = determine_case_phases(case_type, mechanism, is_mascal)
    phase_instr = "This case does NOT require surgery. Only DCR and PCC. Set dcs to null." if phases == ["DCR", "PCC"] else "This case requires surgery. Include DCR, DCS, and PCC."

    zap = str(random.randint(10000, 99999))

    prompt = (
        f"CONTEXT: Role 2 in {environment}, {region}.\n"
        f"CASE: {case_type}\n"
        f"MECHANISM: {mechanism}\n"
        f"{phase_instr}\n"
        f"ZAP NUMBER: Use exactly '{zap}' as the zap field — do not change it.\n"
        f"Generate the case now."
    )

    response = get_client().models.generate_content(
        model=GEMINI_MODEL,
        contents=prompt,
        config={"system_instruction": CASE_SYSTEM_PROMPT, "response_mime_type": "application/json"}
    )

    def _enforce_zap(data: Dict) -> Dict:
        if "zmist" in data:
            data["zmist"]["zap"] = zap
        return data

    try:
        return _enforce_zap(json.loads(response.text))
    except json.JSONDecodeError:
        text = response.text
        start, end = text.find('{'), text.rfind('}') + 1
        if start != -1 and end > start:
            return _enforce_zap(json.loads(text[start:end]))
        raise

def create_fallback_case(case_type: str, mechanism: str, is_trauma: bool = True) -> Dict:
    return {
        "meta": {"title": case_type, "estimated_duration": "30-45 min" if is_trauma else "20-30 min", "personnel": "Medical Team", "target_specialty": "Emergency Medicine" if is_trauma else "Family Physician"},
        "learning_objectives": ["Perform primary survey", "Initiate resuscitation", "Determine evacuation priority"],
        "zmist": {"zap": str(random.randint(10000, 99999)), "mechanism": mechanism, "injuries": case_type, "signs": "HR 110, BP 100/70" if is_trauma else "HR 88, BP 120/80", "treatment": "IV, O2, monitoring"},
        "nine_line": {"line1_location": "Grid TBD", "line2_freq": "Pri: 123.45", "line3_patients_precedence": "1 Alpha" if is_trauma else "1 Charlie", "line4_equipment": "None", "line5_patients_type": "1 Litter" if is_trauma else "1 Ambulatory", "line6_security": "Secure", "line7_marking": "VS-17", "line8_nationality": "US Military", "line9_nbc_terrain": "None"},
        "patient_data": {"demographics": f"{random.randint(19, 35)} yo male Marine", "history": "No PMH", "allergies": "NKDA"},
        "triage_category": "T2" if is_trauma else "T3",
        "disposition": "Evac to Role 3" if is_trauma else "RTD",
        "phases": {"dcr": {"title": "Damage Control Resuscitation", "narrative": "Patient arrives...", "expected_actions": ["Primary survey", "IV access"], "vitals_trend": [{"time": "0:00", "hr": "110", "bp": "100/70", "rr": "20", "spo2": "95%", "gcs": "14"}], "contingencies": [{"condition": "BP <90", "consequence": "Shock", "intervention": "Blood products"}]}, "dcs": None, "pcc": {"title": "Prolonged Casualty Care", "narrative": "Post-resuscitation...", "expected_actions": ["Monitor", "Pain management"], "vitals_trend": [], "contingencies": []}},
        "labs": {"hgb": "11.2", "ph": "7.32", "lactate": "3.1", "base_excess": "-4", "inr": "1.1"},
        "evacuation": {"transport_type": "MEDEVAC", "priority": "Priority" if is_trauma else "Routine", "considerations": "Monitor vitals", "handover_notes": f"Stable s/p {case_type}"},
        "debrief_questions": ["What were key findings?", "Was resuscitation adequate?", "What would you change?"]
    }

# Document generation
def _day_patients(day) -> int:
    """Effective casualties for a day: routine load plus the MASCAL surge."""
    return day.total_patients + ((day.mascal_patients or 0) if day.mascal else 0)

def _day_waves(day) -> int:
    """Effective wave count: routine waves plus one MASCAL wave when enabled."""
    return day.total_waves + (1 if day.mascal else 0)

def generate_warno(config: ExerciseConfig) -> str:
    prompt = f"""Generate USMC WARNO (5-paragraph order) for:
Exercise: {config.exercise_name}, Duration: {config.duration} days
Unit: {config.supported_unit}, Environment: {config.environment}, Region: {config.region}
Threat: {config.threat_level}, Footprint: {', '.join(config.selected_footprint)}
Days: {'; '.join([f"Day {d.day_number}: {d.tactical_setting}, {_day_patients(d)} cas, {'MASCAL '+d.mascal_etiology if d.mascal else 'No MASCAL'}, {'CBRN' if d.cbrn else ''}, {'Night' if d.night_ops else 'Day'}" for d in config.days])}

Include: 1.SITUATION 2.MISSION 3.EXECUTION 4.ADMIN/LOG 5.CMD/SIG"""
    
    response = get_client().models.generate_content(model=GEMINI_MODEL, contents=prompt)
    return response.text

def generate_annex_q(config: ExerciseConfig) -> str:
    total_cas = sum(_day_patients(d) for d in config.days)
    prompt = f"""Generate Annex Q (Medical Services) for:
Exercise: {config.exercise_name}, Environment: {config.environment}, Region: {config.region}
Total Casualties: {total_cas}, Footprint: {', '.join(config.selected_footprint)}
Specialists: {json.dumps(config.specialists)}

Include: HSS concept, MTFs, MEDEVAC, Class VIII, blood support, dental, combat stress, PVNTMED"""
    
    response = get_client().models.generate_content(model=GEMINI_MODEL, contents=prompt)
    return response.text

def generate_medroe(config: ExerciseConfig) -> str:
    has_cbrn = any(d.cbrn for d in config.days)
    has_detainee = any(d.detainee_ops for d in config.days)
    
    prompt = f"""Generate MEDROE for:
Exercise: {config.exercise_name}, Environment: {config.environment}
CBRN: {has_cbrn}, Detainee Ops: {has_detainee}

Include: Treatment priorities, evacuation priorities, holding policy, blood products, documentation, MASCAL procedures"""
    
    response = get_client().models.generate_content(model=GEMINI_MODEL, contents=prompt)
    return response.text

def _road_to_war_fallback(config: "ExerciseConfig") -> str:
    """Templated Road to War video prompt used when the AI call is unavailable."""
    total_cas = sum(_day_patients(d) for d in config.days)
    settings = "; ".join(
        f"Day {d.day_number}: {d.tactical_setting}"
        + (f" — MASCAL ({d.mascal_etiology})" if d.mascal else "")
        for d in config.days
    )
    return f"""ROAD TO WAR — VIDEO GENERATION PROMPT
Exercise: {config.exercise_name}

Paste the block below into Claude to generate an animated "Road to War"
scenario-setting video for this exercise. Edit any bracketed notes first.

------------------------------------------------------------------------
Build a single, self-contained ANIMATED HTML file (inline CSS + JS, no external
files or network requests) that plays as a 90-120 second cinematic "Road to War"
video for the training exercise "{config.exercise_name}". It briefs participants
on the geopolitical backstory leading up to the medical/HSS mission they are
about to rehearse. This is animated HTML — NOT footage from a video-generation
model and no stock video.

SCENARIO FACTS (ground every scene in these):
- Operation: {config.exercise_name}
- Region / theater: {config.region or "[region]"}
- Environment: {config.environment}
- Supported unit: {config.supported_unit}
- Threat level: {config.threat_level}
- Medical footprint deployed: {", ".join(config.selected_footprint) or "[footprint]"}
- Projected casualty load: {total_cas} patients across {len(config.days)} day(s)
- Daily flashpoints: {settings}

HOW TO BUILD IT:
- Animate with CSS keyframes/transitions and animated inline SVG: kinetic
  typography, animated title cards, an animated map (forces massing, sea/air
  lanes closing, weapon-engagement rings), and simple motion graphics.
- VOICEOVER: use the browser Web Speech API (window.speechSynthesis) to read the
  narration aloud, timing the on-screen scenes to the speech. Include a
  full-screen "Begin brief" start button (browsers block autoplay audio) that
  starts both the voiceover and the animation. Show the narration as on-screen
  captions too.
- Desaturated cinematic palette; per-scene durations summing to ~90-120s.
- Optional background music only via the Web Audio API (no external audio).

For EVERY scene, provide the complete word-for-word NARRATION (voiceover) the
speechSynthesis engine will read — the actual lines, not a description.

STRUCTURE:
1. COLD OPEN (~0:00-0:10) — a tense, quiet establishing beat of the {config.environment}
   theater; animated title card "{config.exercise_name}". Include narration.
2. GEOPOLITICAL BACKSTORY (~0:10-0:40) — how tensions in {config.region or "the region"}
   escalated toward conflict; animated map showing forces massing. Include narration.
3. THE FLASHPOINT (~0:40-1:00) — the triggering event that commits the
   {config.supported_unit}; rising tempo, a {config.threat_level} adversary. Include narration.
4. THE MEDICAL PICTURE (~1:00-1:30) — pivot to the HSS challenge: casualty
   flow, MEDEVAC, and the Role 2 mission the audience must be ready for.
   Draw specifics from the exercise's Annex Q (medical concept of support). Include narration.
5. CALL TO ACTION (~1:30-end) — "This is the fight you are training for." Include narration.

MANDATORY: one self-contained animated HTML file with a spoken voiceover and a
full word-for-word narration script (every scene has its spoken lines). Keep it
unclassified and fictional — this is for a training exercise.
------------------------------------------------------------------------
"""

def generate_road_to_war_prompt(config: "ExerciseConfig", annex_text: str = "") -> str:
    """Produce a ready-to-use prompt the user can hand to Claude to generate a
    'Road to War' video tailored to this exercise's Annex Q."""
    total_cas = sum(_day_patients(d) for d in config.days)
    settings = "; ".join(
        f"Day {d.day_number}: {d.tactical_setting}"
        + (f", MASCAL: {d.mascal_etiology}" if d.mascal else "")
        + (", CBRN" if d.cbrn else "")
        for d in config.days
    )
    annex_excerpt = (annex_text or "")[:2500]
    prompt = f"""You are writing a prompt that another person will paste into Claude to
generate a 90-120 second "Road to War" scenario-setting video for a USMC medical
(HSS) training exercise. The video will be produced BY Claude as a single,
self-contained ANIMATED HTML file with a spoken voiceover — NOT footage from a
video-generation model. Write ONLY that prompt — do not write the video/HTML
itself, and do not add commentary before or after.

The prompt you write must:
- Instruct Claude to output ONE self-contained .html file (inline CSS + JS, no
  external files, no network/CDN requests) that plays as an animated Road to War
  video: it establishes the geopolitical backstory escalating to conflict, then
  pivots to the medical threat picture the trainees must be ready for.
- Require the animation to be done with CSS keyframes/transitions and animated
  inline SVG — kinetic typography, animated title cards, an animated map (forces
  massing, sea/air lanes, weapon-engagement rings), and simple motion graphics.
  There is NO live-action footage; do not reference stock video or a video model.
- Require a SPOKEN VOICEOVER using the browser Web Speech API
  (window.speechSynthesis) that reads the narration aloud, advancing/timing the
  on-screen scenes to match the speech. Because browsers block autoplay audio,
  require a full-screen "Begin brief" start button that kicks off both the
  voiceover and the animation on click. Also render the narration as on-screen
  captions/subtitles for accessibility.
- REQUIRE a complete, word-for-word NARRATION (voiceover) script for EVERY
  scene — the actual lines the speechSynthesis engine will read, not a summary.
  This is mandatory: no scene may be left without its spoken narration text.
- Also specify: the scene-by-scene structure (cold open, backstory, flashpoint,
  medical picture, call to action) with per-scene durations that sum to
  ~90-120s, on-screen text/title cards, the desaturated cinematic palette, and
  pacing. (Background music is optional and, if used, must be generated in-page
  via the Web Audio API — no external audio files.)
- Draw the medical picture (casualty flow, MEDEVAC, Role 2 capability) from
  the Annex Q summary provided.
- State that the content is unclassified, fictional, and for training only.

EXERCISE FACTS:
- Operation: {config.exercise_name}
- Region/theater: {config.region}
- Environment: {config.environment}
- Supported unit: {config.supported_unit}
- Threat level: {config.threat_level}
- Medical footprint: {", ".join(config.selected_footprint)}
- Total casualties: {total_cas} across {len(config.days)} day(s)
- Daily settings: {settings}

ANNEX Q (MEDICAL CONCEPT OF SUPPORT) SUMMARY:
{annex_excerpt}

Now write the Road to War video-generation prompt."""
    try:
        response = get_client().models.generate_content(
            model=GEMINI_MODEL, contents=prompt
        )
        text = _response_text(response)
        if text:
            header = f"ROAD TO WAR — VIDEO GENERATION PROMPT\nExercise: {config.exercise_name}\n\nPaste the prompt below into Claude to generate the Road to War video.\n\n"
            return header + text
    except HTTPException:
        raise
    except Exception as e:
        print(f"WARNING: road-to-war prompt AI call failed, using fallback: {e}")
    return _road_to_war_fallback(config)

# Schedule generation
def assign_evaluator(case: Dict, specialists: Dict[str, int], assigned_counts: Dict[str, int]) -> str:
    needs_surgery = case.get("phases", {}).get("dcs") is not None
    triage = case.get("triage_category", "T2")
    
    if needs_surgery:
        priority = ["General Surgery", "Orthopaedic Surgery", "Anesthesiology", "Emergency Medicine"]
    elif triage in ["T1", "T2"]:
        priority = ["Emergency Medicine", "Family Physician", "ER Nurse", "ERC Nurse"]
    else:
        priority = ["ICU Nurse", "Med Surg Nurse", "ER Nurse", "Family Physician"]
    
    abbrev = {"General Surgery": "Gen Surg", "Orthopaedic Surgery": "Ortho", "Emergency Medicine": "EM", "Family Physician": "FP", "Anesthesiology": "Anes", "ERC Nurse": "ERC", "ER Nurse": "ER RN", "ICU Nurse": "ICU RN", "Med Surg Nurse": "MS RN"}
    
    for spec in priority:
        if specialists.get(spec, 0) > 0:
            assigned_counts[spec] = assigned_counts.get(spec, 0) + 1
            num = ((assigned_counts[spec] - 1) % specialists[spec]) + 1
            return f"{abbrev.get(spec, spec)} {num}"
    return "Unassigned"

def _clock(total_minutes: int) -> str:
    """Minutes-from-midnight -> HHMM, wrapping across midnight."""
    total_minutes %= 24 * 60
    return f"{total_minutes // 60:02d}{total_minutes % 60:02d}"

# Evidence-based blood planning. Casualties who require transfusion consume, on
# average, ~6 units whole-blood-equivalent — a mean, not a per-case guess.
# Sources: "The thin red line: Blood planning factors and the enduring need for
# a robust military blood system to support combat operations," J Trauma Acute
# Care Surg 2024 (PMID 38996415), which frames blood-per-casualty planning
# factors; combat DCR/massive-transfusion literature (massive transfusion =
# >=10 RBC units/24h). Only clinically hemorrhaging casualties draw blood.
BLOOD_UNITS_PER_TRANSFUSED = 6
_HEMORRHAGE_KW = ("hemorrhage", "haemorrhage", "amputation", "gsw", "gunshot",
                  "blast", "exsanguination", "massive", "junctional", "vascular",
                  "penetrating", "laceration", "evisceration")

def _requires_blood(case: Dict) -> bool:
    """Derived from the case, not rolled: a casualty needs blood if it goes to
    surgery, is T1, or its injuries/treatment describe hemorrhage."""
    if case.get("phases", {}).get("dcs") is not None:
        return True
    if case.get("triage_category") == "T1":
        return True
    z = case.get("zmist", {})
    text = f"{z.get('injuries','')} {z.get('treatment','')} {z.get('mechanism','')}".lower()
    return any(k in text for k in _HEMORRHAGE_KW)

def _blood_units(case: Dict) -> int:
    """Whole-blood-equivalent units for this casualty: the evidence-based mean
    if transfusion is indicated, otherwise zero."""
    return BLOOD_UNITS_PER_TRANSFUSED if _requires_blood(case) else 0

def _evac_min(precedence: str) -> int:
    """POI->Role 2 transit planning bracket by evacuation precedence."""
    p = (precedence or "").lower()
    if "urgent" in p:
        return 30
    if "priority" in p:
        return 60
    return 120  # routine / convenience

def _dwell_min(case: Dict) -> int:
    """Role 2 treatment/OR occupancy bracket, used for the saturation picture."""
    phases = case.get("phases", {}) or {}
    if phases.get("dcs"):
        return 150  # surgical: OR + immediate post-op hold
    if phases.get("pcc"):
        return 90   # resus + prolonged care hold
    return 45

def _route_and_inbound(case: Dict, is_mascal_wave: bool) -> tuple:
    """Decide route + whether the casualty is called in to the COC, using only
    facts already in the case (triage, surgical need, stated transport). No dice:
    minor DNBI (T3/T4, non-surgical) self-present as walking wounded; everyone
    else — and all MASCAL casualties — are inbound with a COC call + 9-line."""
    triage = case.get("triage_category", "T2")
    surgical = case.get("phases", {}).get("dcs") is not None
    transport = (case.get("evacuation", {}).get("transport_type") or "").strip()
    inbound = is_mascal_wave or surgical or triage in ("T1", "T2")
    if not inbound:
        return "Walk-in", False
    if transport and not any(w in transport.lower() for w in ("walk", "ambul", "self")):
        return transport, True
    return ("Litter" if is_mascal_wave else "MEDEVAC"), True

def _case_teo(case: Dict) -> Dict:
    """Flatten the case's own clinical detail for the T&EO sheet. Pulls straight
    from the generated case — nothing is invented here."""
    phases = case.get("phases", {}) or {}
    def _collect(field):
        out = []
        for pk in ("dcr", "dcs", "pcc"):
            ph = phases.get(pk)
            if isinstance(ph, dict):
                out += ph.get(field, []) or []
        return out
    care_level = "DCR" + ("+DCS" if phases.get("dcs") else "") + ("+PCC" if phases.get("pcc") else "")
    contingencies = [c for c in _collect("contingencies") if isinstance(c, dict)]
    evac = case.get("evacuation", {}) or {}
    return {
        "zap": case.get("zmist", {}).get("zap", ""),
        "surgical": "Yes" if phases.get("dcs") else "No",
        "care_level": care_level,
        "disposition": case.get("disposition", ""),
        "blood_units": _blood_units(case),
        "r2_dwell": _dwell_min(case),
        "evac_precedence": evac.get("priority", ""),
        "onward": evac.get("transport_type", ""),
        "signs": case.get("zmist", {}).get("signs", ""),
        "handover": evac.get("handover_notes", ""),
        "expected": " • ".join(str(a) for a in _collect("expected_actions")),
        "contingencies": " | ".join(
            f"IF {c.get('condition','')} → {c.get('consequence','')} → {c.get('intervention','')}"
            for c in contingencies
        ),
        "debrief": " • ".join(case.get("debrief_questions", []) or []),
    }

def generate_schedule(config: ExerciseConfig, cases: List[Dict]) -> List[Dict]:
    schedule = []
    case_idx = 0
    assigned_counts = {}

    for day in config.days:
        if day.cbrn:
            schedule.append({"day": day.day_number, "event": "DRILL", "coc_hit_time": "N/A", "time": "0900", "nine_line_time": "N/A", "route": "N/A", "triage_cat": "N/A", "mechanism": "CBRN DRILL", "brief_description": "1-hour CBRN exercise - All clinical ops paused", "evaluator": "All Hands", "case_num": "DRILL"})
        if day.detainee_ops:
            schedule.append({"day": day.day_number, "event": "INJECT", "coc_hit_time": "N/A", "time": "1300", "nine_line_time": "N/A", "route": "N/A", "triage_cat": "N/A", "mechanism": "DETAINEE / EPW", "brief_description": "Detainee/EPW presents for care — apply detainee handling + MEDROE", "evaluator": "All Hands", "case_num": "INJECT"})

        # Night ops means SOME casualties arrive at night, not all care at night:
        # carve a share of the routine load into a dedicated night wave.
        night_pts = min(max(1, round(day.total_patients * 0.35)), day.total_patients) if day.night_ops else 0
        day_pts = day.total_patients - night_pts

        # Routine day waves carry the daytime load; a MASCAL adds one more wave
        # on top with its own surge count (additive, not an override).
        base_waves = max(day.total_waves, 1)
        pts_per_wave = day_pts // base_waves
        remainder = day_pts % base_waves
        day_waves = [{"pts": pts_per_wave + (1 if w < remainder else 0), "mascal": False}
                     for w in range(base_waves)]
        if day.mascal and day.mascal_patients:
            day_waves.append({"pts": day.mascal_patients, "mascal": True})
        night_waves = [{"pts": night_pts, "mascal": False}] if night_pts else []

        # Day window 0700-1900 (12h); night window 2000-0200 (6h).
        for group, start_min, span_min in ((day_waves, 7 * 60, 12 * 60), (night_waves, 20 * 60, 6 * 60)):
            interval = span_min / (len(group) + 1) if group else 0
            for k, w in enumerate(group):
                w["start_min"] = start_min + int(interval * (k + 1))

        # Emit chronologically: day waves, then the night wave.
        for w in day_waves + night_waves:
            wave_pts = w["pts"]
            is_mascal_wave = w["mascal"]
            time_spread = 45 if is_mascal_wave else 60

            for p in range(wave_pts):
                if case_idx >= len(cases):
                    break
                case = cases[case_idx]

                # MASCAL is a compressed spike then a tail (front-loaded); routine
                # casualties spread evenly across the wave.
                frac = p / max(wave_pts, 1)
                offset = int((frac ** 2) * 40) if is_mascal_wave else int(frac * 60)
                arr_total = w["start_min"] + offset

                # Route + COC call are derived from the case, not rolled at random.
                route, inbound = _route_and_inbound(case, is_mascal_wave)
                if inbound:
                    coc_hit_time = _clock(arr_total - 45)
                    nine_time = _clock(arr_total - 30)
                else:
                    coc_hit_time = "N/A"
                    nine_time = "N/A"

                if is_mascal_wave:
                    event = f"MASCAL ({day.mascal_etiology})" if day.mascal_etiology else "MASCAL"
                else:
                    event = "Routine"

                teo = _case_teo(case)
                # Timeline: point of injury (golden-hour anchor) back from arrival
                # by the precedence transit bracket; Role 2 cleared = arrival + dwell.
                poi_time = _clock(arr_total - _evac_min(teo["evac_precedence"]))
                cleared = _clock(arr_total + teo["r2_dwell"])
                schedule.append({
                    "day": day.day_number,
                    "event": event,
                    "poi_time": poi_time,
                    "cleared": cleared,
                    "coc_hit_time": coc_hit_time,
                    "time": _clock(arr_total),
                    "nine_line_time": nine_time,
                    "route": route,
                    "triage_cat": case.get("triage_category", "T2"),
                    "mechanism": case.get("zmist", {}).get("mechanism", "Unknown")[:50],
                    "brief_description": case.get("zmist", {}).get("injuries", "")[:80],
                    "evaluator": assign_evaluator(case, config.specialists, assigned_counts),
                    "case_num": f"Case {case_idx + 1}",
                    **teo,
                })
                case_idx += 1

    return schedule

# Document creation
def create_docx(title: str, subtitle: str, content: str) -> BytesIO:
    doc = Document()
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph(subtitle)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].bold = True
    doc.add_paragraph(f'DTG: {datetime.now().strftime("%d%H%MZ %b %Y").upper()}')
    doc.add_paragraph('CLASSIFICATION: UNCLASSIFIED // FOR TRAINING ONLY')
    doc.add_paragraph()
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            if any(line.strip().startswith(f'{i}.') for i in range(1, 10)) or line.strip().isupper():
                if p.runs:
                    p.runs[0].bold = True
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def create_case_book(cases: List[Dict], config: ExerciseConfig) -> BytesIO:
    doc = Document()
    
    # Title
    h = doc.add_heading(config.exercise_name.upper(), 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph('SIMULATION CASE BOOK')
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].bold = True
    s.runs[0].font.size = Pt(18)
    doc.add_paragraph()
    doc.add_paragraph(f'Environment: {config.environment} | Region: {config.region}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Duration: {config.duration} days | Total Cases: {len(cases)}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%d %b %Y %H%M")}')
    doc.add_page_break()
    
    # TOC
    doc.add_heading('TABLE OF CONTENTS', level=1)
    for i, case in enumerate(cases):
        doc.add_paragraph(f'Case {i+1}: {case.get("meta", {}).get("title", "Untitled")} (ZAP: {case.get("zmist", {}).get("zap", "00000")})')
    doc.add_page_break()
    
    # Cases
    for i, case in enumerate(cases):
        meta = case.get("meta", {})
        doc.add_heading(f'CASE {i+1}: {meta.get("title", "Untitled")}', level=1)
        
        # Meta
        t = doc.add_table(rows=2, cols=4)
        t.style = 'Table Grid'
        t.rows[0].cells[0].text = 'Duration'
        t.rows[0].cells[1].text = meta.get("estimated_duration", "N/A")
        t.rows[0].cells[2].text = 'Personnel'
        t.rows[0].cells[3].text = meta.get("personnel", "N/A")
        t.rows[1].cells[0].text = 'Specialty'
        t.rows[1].cells[1].text = meta.get("target_specialty", "N/A")
        t.rows[1].cells[2].text = 'Triage'
        t.rows[1].cells[3].text = case.get("triage_category", "N/A")
        doc.add_paragraph()
        
        # Learning Objectives
        doc.add_heading('Learning Objectives', level=2)
        for obj in case.get("learning_objectives", []):
            doc.add_paragraph(f'• {obj}')
        
        # Z-MIST
        doc.add_heading('Z-MIST Report', level=2)
        zmist = case.get("zmist", {})
        zt = doc.add_table(rows=5, cols=2)
        zt.style = 'Table Grid'
        for idx, (lbl, key) in enumerate([('Z - Zap Number', 'zap'), ('M - Mechanism', 'mechanism'), ('I - Injuries', 'injuries'), ('S - Signs', 'signs'), ('T - Treatment', 'treatment')]):
            zt.rows[idx].cells[0].text = lbl
            zt.rows[idx].cells[1].text = str(zmist.get(key, ""))
        doc.add_paragraph()
        
        # 9-Line
        doc.add_heading('9-Line MEDEVAC Request', level=2)
        nl = case.get("nine_line", {})
        nt = doc.add_table(rows=9, cols=2)
        nt.style = 'Table Grid'
        for idx, (lbl, key) in enumerate([('Line 1 - Location', 'line1_location'), ('Line 2 - Frequency', 'line2_freq'), ('Line 3 - Patients/Precedence', 'line3_patients_precedence'), ('Line 4 - Equipment', 'line4_equipment'), ('Line 5 - Patient Type', 'line5_patients_type'), ('Line 6 - Security', 'line6_security'), ('Line 7 - Marking', 'line7_marking'), ('Line 8 - Nationality', 'line8_nationality'), ('Line 9 - NBC/Terrain', 'line9_nbc_terrain')]):
            nt.rows[idx].cells[0].text = lbl
            nt.rows[idx].cells[1].text = str(nl.get(key, ""))
        doc.add_paragraph()
        
        # Patient
        doc.add_heading('Patient Information', level=2)
        pt = case.get("patient_data", {})
        doc.add_paragraph(f'Demographics: {pt.get("demographics", "N/A")}')
        doc.add_paragraph(f'Medical History: {pt.get("history", "N/A")}')
        doc.add_paragraph(f'Allergies: {pt.get("allergies", "NKDA")}')
        
        # Labs
        labs = case.get("labs", {})
        if labs and any(labs.values()):
            doc.add_heading('Initial Labs', level=2)
            lt = doc.add_table(rows=2, cols=5)
            lt.style = 'Table Grid'
            for c, h in enumerate(['Hgb', 'pH', 'Lactate', 'Base Excess', 'INR']):
                lt.rows[0].cells[c].text = h
            for c, k in enumerate(['hgb', 'ph', 'lactate', 'base_excess', 'inr']):
                lt.rows[1].cells[c].text = str(labs.get(k, ""))
        doc.add_paragraph()
        
        # Phases
        phases = case.get("phases", {})
        for pk in ["dcr", "dcs", "pcc"]:
            phase = phases.get(pk)
            if not phase:
                continue
            doc.add_heading(phase.get("title", pk.upper()), level=2)
            doc.add_paragraph(phase.get("narrative", ""))
            
            doc.add_paragraph('Expected Actions:')
            for act in phase.get("expected_actions", []):
                doc.add_paragraph(f'☐ {act}')
            
            vitals = phase.get("vitals_trend", [])
            if vitals:
                doc.add_paragraph()
                doc.add_paragraph('Vitals Trend:')
                vt = doc.add_table(rows=len(vitals)+1, cols=6)
                vt.style = 'Table Grid'
                for c, h in enumerate(['Time', 'HR', 'BP', 'RR', 'SpO2', 'GCS']):
                    vt.rows[0].cells[c].text = h
                for r, v in enumerate(vitals):
                    for c, k in enumerate(['time', 'hr', 'bp', 'rr', 'spo2', 'gcs']):
                        vt.rows[r+1].cells[c].text = str(v.get(k, ""))
            
            cont = phase.get("contingencies", [])
            if cont:
                doc.add_paragraph()
                doc.add_paragraph('Contingencies (If/Then):')
                for c in cont:
                    doc.add_paragraph(f'IF: {c.get("condition", "")}')
                    doc.add_paragraph(f'   → CONSEQUENCE: {c.get("consequence", "")}')
                    doc.add_paragraph(f'   → INTERVENTION: {c.get("intervention", "")}')
            doc.add_paragraph()
        
        # Evacuation
        doc.add_heading('Evacuation / En Route Care', level=2)
        evac = case.get("evacuation", {})
        doc.add_paragraph(f'Transport: {evac.get("transport_type", "N/A")}')
        doc.add_paragraph(f'Priority: {evac.get("priority", "N/A")}')
        doc.add_paragraph(f'Considerations: {evac.get("considerations", "")}')
        doc.add_paragraph(f'Handover: {evac.get("handover_notes", "")}')
        
        # Debrief
        doc.add_heading('Debrief Questions', level=2)
        for q in case.get("debrief_questions", []):
            doc.add_paragraph(f'• {q}')
        
        if i < len(cases) - 1:
            doc.add_page_break()
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def _autosize(ws, df):
    for idx, col in enumerate(df.columns):
        body = df[col].astype(str).map(len).max() if len(df) else 0
        max_len = max(int(body or 0), len(str(col))) + 2
        ws.column_dimensions[chr(65 + idx)].width = min(max_len, 60)

def _sheet(df: pd.DataFrame, source_cols, labels) -> pd.DataFrame:
    """Select/relabel columns, filling any missing ones (older stored exercises
    predate some keys) so re-download never breaks."""
    for c in source_cols:
        if c not in df.columns:
            df[c] = ""
    out = df[source_cols].copy()
    # Inject rows (CBRN/detainee) lack the clinical keys -> NaN; blank them.
    out = out.fillna("")
    out.columns = labels
    return out

def create_msel(schedule: List[Dict], config: ExerciseConfig) -> BytesIO:
    raw = pd.DataFrame(schedule)

    # Sheet 1 — MSEL timeline (what EXCON runs the exercise from).
    msel = _sheet(
        raw,
        ['day', 'event', 'poi_time', 'coc_hit_time', 'nine_line_time', 'time', 'route', 'evac_precedence', 'triage_cat', 'surgical', 'disposition', 'zap', 'mechanism', 'brief_description', 'evaluator', 'case_num'],
        ['Day', 'Event', 'POI Time', 'COC Hit Time', '9-Line Time', 'Arrival', 'Route', 'Precedence', 'Triage', 'Surgical', 'Disposition', 'ZAP #', 'Mechanism', 'Description', 'Evaluator', 'Serial'],
    )

    # Sheet 2 — T&EO / Controller detail (per-casualty, straight from the case).
    teo = _sheet(
        raw,
        ['day', 'time', 'cleared', 'zap', 'triage_cat', 'care_level', 'surgical', 'blood_units', 'r2_dwell', 'disposition', 'signs', 'onward', 'handover', 'expected', 'contingencies', 'debrief', 'evaluator'],
        ['Day', 'Arrival', 'R2 Cleared', 'ZAP #', 'Triage', 'Care Level', 'Surgical', 'Blood (WBE u)', 'R2 Dwell (min)', 'Disposition', 'Initial Signs', 'Onward Tpt', 'Handover', 'Expected Key Actions', 'Contingencies (If/Then)', 'Debrief Prompts', 'Evaluator'],
    )

    # Sheet 3 — Objectives (exercise METL / footprint / personnel, from config).
    obj_data = []
    for m in config.selected_mets:
        obj_data.append({"Category": "METL Task", "Item": m})
    for f in config.selected_footprint:
        obj_data.append({"Category": "Footprint", "Item": f})
    for k, v in (config.specialists or {}).items():
        if v:
            obj_data.append({"Category": "Personnel", "Item": f"{k}: {v}"})
    # Aggregate blood-demand planning factor (whole-blood-equivalent) — sum of
    # the evidence-based per-transfused figure across casualties that draw blood.
    total_blood = int(pd.to_numeric(raw.get("blood_units"), errors="coerce").fillna(0).sum()) if "blood_units" in raw.columns else 0
    transfused = int((pd.to_numeric(raw.get("blood_units"), errors="coerce").fillna(0) > 0).sum()) if "blood_units" in raw.columns else 0
    obj_data.append({"Category": "Planning Factor", "Item": f"Blood demand: {total_blood} WBE units ({transfused} casualties transfused @ {BLOOD_UNITS_PER_TRANSFUSED} u avg)"})
    objectives = pd.DataFrame(obj_data) if obj_data else pd.DataFrame({"Category": [], "Item": []})

    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for name, d in (('MSEL', msel), ('T&EO', teo), ('Objectives', objectives)):
            d.to_excel(writer, index=False, sheet_name=name)
            _autosize(writer.sheets[name], d)
    output.seek(0)
    return output

# API Endpoints
@app.get("/")
async def root():
    return {"status": "Role 2 Exercise Builder API", "version": "2.0"}

@app.get("/health")
async def health():
    return {"status": "healthy"}

@app.post("/generate-name")
async def generate_name_endpoint(data: dict):
    env = data.get("environment", "General")
    region = data.get("region", "")
    threat = data.get("threatLevel", "")
    unit = data.get("supportedUnit", "Medical")
    
    prompt = f"""Generate one military exercise name for a USMC {unit} HSS training exercise in {env} terrain, {region}, facing a {threat} threat.

Format: "Operation [Descriptor] [Noun]" — 2 words after Operation, or occasionally 3 if it flows well.

Structure:
- Descriptor: an adjective that evokes the operational environment or tempo (e.g., Silent, Relentless, Fractured, Bleeding, Stalking, Burning, Severed, Hollow, Bitter)
- Noun: pick ONE category — fierce animal (Viper, Jackal, Condor, Wolverine, Mako, Lynx, Mantis), masculine emotion (Fury, Wrath, Valor, Resolve, Defiance, Reckoning, Anguish), or military/defense concept (Rampart, Bastion, Crucible, Gauntlet, Threshold, Salient, Bulwark)

Banned words: Iron, Steel, Crimson, Eagle, Thunder, Guardian, Shield, Ghost, Phantom, Storm, Black, Red, Blue, Dragon, Tiger, Wolf, Hawk, Viper (if region is not maritime)

Return ONLY the operation name. Nothing else. Seed: {random.random()}"""
    # The violent descriptors in this prompt can trip Gemini's safety filters,
    # which makes response.text raise and the endpoint 500 ("not fetching").
    # Relax the filters where allowed, extract text safely, and fall back to a
    # curated local name so the button always returns a usable result.
    name = ""
    try:
        try:
            from google.genai import types
            gen_config = types.GenerateContentConfig(
                temperature=1.3,
                safety_settings=[
                    types.SafetySetting(category=c, threshold="BLOCK_NONE")
                    for c in (
                        "HARM_CATEGORY_HARASSMENT",
                        "HARM_CATEGORY_HATE_SPEECH",
                        "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                        "HARM_CATEGORY_DANGEROUS_CONTENT",
                    )
                ],
            )
            response = get_client().models.generate_content(
                model=GEMINI_MODEL, contents=prompt, config=gen_config
            )
        except Exception:
            # Older/newer SDKs may not accept the config shape above — retry plain.
            response = get_client().models.generate_content(
                model=GEMINI_MODEL, contents=prompt
            )
        name = _response_text(response).replace('"', '').replace("'", "")
        # Models sometimes return a short explanation before the name.
        name = name.splitlines()[0].strip() if name else ""
    except HTTPException:
        raise  # e.g. GEMINI_API_KEY missing — surface it as-is
    except Exception as e:
        print(f"WARNING: generate-name AI call failed, using local fallback: {e}")

    if not name:
        name = _local_exercise_name(region)
    if not name.startswith("Operation"):
        name = f"Operation {name}"
    return {"name": name}

def _build_case_tasks(config: ExerciseConfig) -> List[tuple]:
    """Return a list of (case_type, mechanism, is_trauma, is_mascal) tuples — one per patient."""
    tasks = []
    env_dnbi = DNBI_BY_ENVIRONMENT.get(config.environment, []) + DNBI_BY_ENVIRONMENT.get("General", [])
    for day in config.days:
        # Routine casualty load for the day (not part of any MASCAL surge).
        routine_ratio = 0.85 if day.tactical_setting in ["Frontal Attack", "Amphibious Assault"] else 0.50
        num_trauma = int(day.total_patients * routine_ratio)
        num_dnbi = day.total_patients - num_trauma
        for _ in range(num_trauma):
            tasks.append((random.choice(GENERAL_TRAUMA), day.tactical_setting, True, False))
        for _ in range(num_dnbi):
            tasks.append((random.choice(env_dnbi), f"DNBI - {config.environment}", False, False))
        # MASCAL surge — additive extra casualties, trauma-heavy, tagged is_mascal.
        if day.mascal and day.mascal_patients:
            inj_types = TRAUMA_BY_ETIOLOGY.get(day.mascal_etiology, GENERAL_TRAUMA) if day.mascal_etiology else GENERAL_TRAUMA
            m_trauma = int(day.mascal_patients * 0.85)
            for _ in range(m_trauma):
                tasks.append((random.choice(inj_types), day.mascal_etiology or day.tactical_setting, True, True))
            for _ in range(day.mascal_patients - m_trauma):
                tasks.append((random.choice(env_dnbi), f"DNBI - {config.environment}", False, True))
    return tasks

def _generate_one(task: tuple, environment: str, region: str) -> Dict:
    case_type, mech, is_trauma, is_mascal = task
    try:
        return generate_case_sync(case_type, mech, environment, region, is_mascal)
    except:
        return create_fallback_case(case_type, mech, is_trauma)

# In-memory job store — always used as primary, DB synced as best-effort
_jobs: Dict[str, Dict] = {}

def _job_create(job_id: str):
    _jobs[job_id] = {"status": "running", "progress": "Starting...", "completed": 0, "total": 0}
    if SessionLocal:
        try:
            db = SessionLocal()
            try:
                db.add(Job(id=job_id))
                db.commit()
            finally:
                db.close()
        except Exception as e:
            print(f"DB job create failed (using memory only): {e}")

def _job_update(job_id: str, **kwargs):
    if job_id in _jobs:
        _jobs[job_id].update(kwargs)
    if SessionLocal:
        try:
            db = SessionLocal()
            try:
                job = db.query(Job).filter(Job.id == job_id).first()
                if job:
                    for k, v in kwargs.items():
                        setattr(job, k, v)
                    db.commit()
            finally:
                db.close()
        except Exception as e:
            print(f"DB job update failed: {e}")

def _job_get(job_id: str):
    # Memory first (fast, same-instance), DB as cross-instance fallback
    if job_id in _jobs:
        return _jobs[job_id]
    if SessionLocal:
        try:
            db = SessionLocal()
            try:
                job = db.query(Job).filter(Job.id == job_id).first()
                if job:
                    return {"status": job.status, "progress": job.progress,
                            "completed": job.completed, "total": job.total,
                            "token": job.token, "filename": job.filename, "error": job.error}
            finally:
                db.close()
        except Exception as e:
            print(f"DB job get failed: {e}")
    return None

def _run_generation(config: ExerciseConfig, job_id: str):
    try:
        tasks = _build_case_tasks(config)
        total = len(tasks)
        _job_update(job_id, progress="Generating cases...", completed=0, total=total)

        cases_results: Dict[int, Dict] = {}
        with ThreadPoolExecutor(max_workers=min(5, max(total, 1))) as pool:
            futures = {pool.submit(_generate_one, t, config.environment, config.region): i
                       for i, t in enumerate(tasks)}
            completed = 0
            for future in as_completed(futures):
                cases_results[futures[future]] = future.result()
                completed += 1
                _job_update(job_id, progress=f"Generating cases: {completed} / {total}", completed=completed)

        cases = [cases_results[i] for i in range(total)]
        _job_update(job_id, progress="Generating documents...", completed=total)

        random.shuffle(cases)
        schedule = generate_schedule(config, cases)
        with ThreadPoolExecutor(max_workers=3) as pool:
            f_warno = pool.submit(generate_warno, config)
            f_annex = pool.submit(generate_annex_q, config)
            f_medroe = pool.submit(generate_medroe, config)
            warno = f_warno.result()
            annex = f_annex.result()
            medroe = f_medroe.result()

        # Road to War video prompt is derived from the freshly generated Annex Q.
        _job_update(job_id, progress="Writing Road to War prompt...")
        road_to_war = generate_road_to_war_prompt(config, annex)

        _job_update(job_id, progress="Assembling package...")

        if SessionLocal:
            db = SessionLocal()
            try:
                ex = Exercise(name=config.exercise_name, config=config.dict(), cases=cases,
                              msel_data=schedule, warno_text=warno, annex_q_text=annex, medroe_text=medroe,
                              road_to_war_text=road_to_war)
                db.add(ex)
                db.commit()
            finally:
                db.close()

        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{config.exercise_name}_MSEL.xlsx", create_msel(schedule, config).getvalue())
            zf.writestr(f"{config.exercise_name}_WARNO.docx", create_docx("WARNING ORDER", config.exercise_name.upper(), warno).getvalue())
            zf.writestr(f"{config.exercise_name}_Annex_Q.docx", create_docx("ANNEX Q (MEDICAL SERVICES)", f"TO OPORD {config.exercise_name.upper()}", annex).getvalue())
            zf.writestr(f"{config.exercise_name}_MEDROE.docx", create_docx("MEDICAL RULES OF ENGAGEMENT", config.exercise_name.upper(), medroe).getvalue())
            zf.writestr(f"{config.exercise_name}_Case_Book.docx", create_case_book(cases, config).getvalue())
            zf.writestr(f"{config.exercise_name}_Road_to_War_Prompt.docx", create_docx("ROAD TO WAR — VIDEO PROMPT", config.exercise_name.upper(), road_to_war).getvalue())
        zip_buf.seek(0)

        token = str(uuid.uuid4())
        _packages[token] = zip_buf.getvalue()
        _job_update(job_id, status="complete", progress="Package ready!",
                    completed=total, total=total, token=token,
                    filename=f"{config.exercise_name}_Package.zip")
    except Exception as e:
        import traceback
        traceback.print_exc()
        _job_update(job_id, status="error", progress=str(e), error=str(e))


@app.post("/generate-exercise")
async def generate_exercise(config: ExerciseConfig):
    job_id = str(uuid.uuid4())
    _job_create(job_id)
    threading.Thread(target=_run_generation, args=(config, job_id), daemon=True).start()
    return {"job_id": job_id}


@app.get("/jobs/{job_id}")
async def get_job_status(job_id: str):
    job = _job_get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@app.get("/download/{token}")
async def download_package(token: str):
    zip_bytes = _packages.pop(token, None)
    if not zip_bytes:
        raise HTTPException(status_code=404, detail="Package not found or already downloaded")
    return Response(zip_bytes, media_type="application/zip",
                    headers={"Content-Disposition": "attachment; filename=package.zip"})

@app.get("/exercises")
async def list_exercises():
    if not SessionLocal:
        return {"exercises": []}
    db = SessionLocal()
    try:
        exs = db.query(Exercise).order_by(Exercise.created_at.desc()).all()
        exercises_list = []
        for e in exs:
            try:
                # Parse JSON if stored as string
                config = e.config or {}
                if isinstance(config, str):
                    config = json.loads(config)
                cases = e.cases or []
                if isinstance(cases, str):
                    cases = json.loads(cases)

                exercises_list.append({
                    "id": e.id,
                    "name": e.name,
                    "created_at": e.created_at.isoformat() if e.created_at else None,
                    "duration": config.get("duration") if isinstance(config, dict) else None,
                    "environment": config.get("environment") if isinstance(config, dict) else None,
                    "total_cases": len(cases) if isinstance(cases, list) else 0
                })
            except Exception as ex:
                # Log but skip malformed exercises
                exercises_list.append({
                    "id": e.id,
                    "name": e.name or "Unknown",
                    "created_at": None,
                    "duration": None,
                    "environment": None,
                    "total_cases": 0
                })
        return {"exercises": exercises_list}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")
    finally:
        db.close()

@app.get("/exercises/{exercise_id}")
async def get_exercise(exercise_id: int):
    if not SessionLocal:
        raise HTTPException(status_code=404, detail="DB not configured")
    db = SessionLocal()
    try:
        ex = db.query(Exercise).filter(Exercise.id == exercise_id).first()
        if not ex:
            raise HTTPException(status_code=404, detail="Not found")
        return {"id": ex.id, "name": ex.name, "created_at": ex.created_at.isoformat() if ex.created_at else None, "config": ex.config, "cases": ex.cases, "msel_data": ex.msel_data}
    finally:
        db.close()

@app.get("/exercises/{exercise_id}/download")
async def download_exercise(exercise_id: int):
    if not SessionLocal:
        raise HTTPException(status_code=404, detail="DB not configured")
    db = SessionLocal()
    try:
        ex = db.query(Exercise).filter(Exercise.id == exercise_id).first()
        if not ex:
            raise HTTPException(status_code=404, detail="Not found")

        # Parse JSON if stored as string
        config_data = ex.config
        if isinstance(config_data, str):
            config_data = json.loads(config_data)
        config = ExerciseConfig(**config_data)

        msel_data = ex.msel_data
        if isinstance(msel_data, str):
            msel_data = json.loads(msel_data)

        cases = ex.cases
        if isinstance(cases, str):
            cases = json.loads(cases)

        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.writestr(f"{config.exercise_name}_MSEL.xlsx", create_msel(msel_data, config).getvalue())
            zf.writestr(f"{config.exercise_name}_WARNO.docx", create_docx("WARNING ORDER", config.exercise_name.upper(), ex.warno_text).getvalue())
            zf.writestr(f"{config.exercise_name}_Annex_Q.docx", create_docx("ANNEX Q (MEDICAL SERVICES)", f"TO OPORD {config.exercise_name.upper()}", ex.annex_q_text).getvalue())
            zf.writestr(f"{config.exercise_name}_MEDROE.docx", create_docx("MEDICAL RULES OF ENGAGEMENT", config.exercise_name.upper(), ex.medroe_text).getvalue())
            zf.writestr(f"{config.exercise_name}_Case_Book.docx", create_case_book(cases, config).getvalue())
            road_to_war = getattr(ex, "road_to_war_text", None) or generate_road_to_war_prompt(config, ex.annex_q_text or "")
            zf.writestr(f"{config.exercise_name}_Road_to_War_Prompt.docx", create_docx("ROAD TO WAR — VIDEO PROMPT", config.exercise_name.upper(), road_to_war).getvalue())
        zip_buf.seek(0)
        return Response(zip_buf.getvalue(), headers={'Content-Disposition': f'attachment; filename="{config.exercise_name}_Package.zip"'}, media_type='application/zip')
    finally:
        db.close()

@app.get("/exercises/{exercise_id}/document/{doc_type}")
async def download_document(exercise_id: int, doc_type: str):
    if not SessionLocal:
        raise HTTPException(status_code=404, detail="DB not configured")
    if doc_type not in ["msel", "warno", "annex_q", "medroe", "case_book", "road_to_war"]:
        raise HTTPException(status_code=400, detail="Invalid doc type")
    
    db = SessionLocal()
    try:
        ex = db.query(Exercise).filter(Exercise.id == exercise_id).first()
        if not ex:
            raise HTTPException(status_code=404, detail="Not found")

        # Parse JSON if stored as string
        config_data = ex.config
        if isinstance(config_data, str):
            config_data = json.loads(config_data)
        config = ExerciseConfig(**config_data)

        msel_data = ex.msel_data
        if isinstance(msel_data, str):
            msel_data = json.loads(msel_data)

        cases = ex.cases
        if isinstance(cases, str):
            cases = json.loads(cases)

        if doc_type == "msel":
            buf, fn, mt = create_msel(msel_data, config), f"{config.exercise_name}_MSEL.xlsx", 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif doc_type == "warno":
            buf, fn, mt = create_docx("WARNING ORDER", config.exercise_name.upper(), ex.warno_text), f"{config.exercise_name}_WARNO.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif doc_type == "annex_q":
            buf, fn, mt = create_docx("ANNEX Q", f"TO OPORD {config.exercise_name.upper()}", ex.annex_q_text), f"{config.exercise_name}_Annex_Q.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif doc_type == "medroe":
            buf, fn, mt = create_docx("MEDROE", config.exercise_name.upper(), ex.medroe_text), f"{config.exercise_name}_MEDROE.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif doc_type == "road_to_war":
            rtw = getattr(ex, "road_to_war_text", None) or generate_road_to_war_prompt(config, ex.annex_q_text or "")
            buf, fn, mt = create_docx("ROAD TO WAR — VIDEO PROMPT", config.exercise_name.upper(), rtw), f"{config.exercise_name}_Road_to_War_Prompt.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            buf, fn, mt = create_case_book(cases, config), f"{config.exercise_name}_Case_Book.docx", 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        return Response(buf.getvalue(), headers={'Content-Disposition': f'attachment; filename="{fn}"'}, media_type=mt)
    finally:
        db.close()